import json
import os
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import openai
import time
import uuid
from dotenv import load_dotenv
from docx import Document

load_dotenv()

app = Flask(__name__)
CORS(app)

# Set your OpenAI API key and assistant ID
openai.api_key = os.getenv("API_KEY")
assistant_id = os.getenv("ASSISTANT_ID")

# Get the directory of the current script
app_dir = os.path.dirname(os.path.abspath(__file__))

def create_thread():
    """Create a thread for conversation."""
    thread = openai.beta.threads.create()
    return thread.id

def submit_message(thread_id, user_message):
    """Submit a message to the thread."""
    openai.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=user_message
    )

def run_assistant(thread_id):
    """Run the assistant and get a response."""
    run = openai.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id,
    )
    return run.id

def check_status(run_id, thread_id):
    """Check the status of the run."""
    run = openai.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)
    return run.status

def get_response(thread_id):
    response = openai.beta.threads.messages.list(thread_id=thread_id, order="desc")

    for msg in response.data:
        if msg.role == "assistant":
            return msg.content[0].text.value

    return None

def get_responseOld(thread_id):
    """Retrieve messages from the thread."""
    response = openai.beta.threads.messages.list(thread_id=thread_id, order="asc")
    return response.data[-1].content[0].text.value if response.data else None

import collections.abc

def fill_template(data, template_path, output_path):
    """Fill the .docx template with provided data and save it as a new file."""
    def replace_placeholders(paragraph, data):
        for key, value in data.items():
            if isinstance(value, collections.abc.Mapping):
                replace_placeholders(paragraph, value)
            else:
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    doc = Document(template_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders(paragraph, data)
                    if "ishodi_ucenja" in data:
                        for i in range(1, 5):
                            iu_key = f"iu_{i}"
                            if iu_key in data["ishodi_ucenja"]:
                                iu = data["ishodi_ucenja"][iu_key]
                                for key, value in iu.items():
                                    placeholder = f"{{{{{iu_key}}}}}"
                                    if placeholder in paragraph.text:
                                        paragraph.text = paragraph.text.replace(placeholder, str(value))

                                iu_kod_key = f"iu_kod_{i}"
                                iu_kod = data["ishodi_ucenja"].get(iu_key, {}).get("kod", "")
                                placeholder = f"{{{{{iu_kod_key}}}}}"
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(iu_kod))

                                iu_spkod_key = f"iu_spkod_{i}"
                                iu_spkod = data["ishodi_ucenja"].get(iu_key, {}).get("sp_kod", "")
                                placeholder = f"{{{{{iu_spkod_key}}}}}"
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(iu_spkod))

                    if "literatura" in data:
                        for i in range(1, 4):
                            literatura_ob_key = f"literatura_ob_{i}"
                            literatura_ob = data["literatura"].get("obavezna", [])[i - 1] if i - 1 < len(data["literatura"].get("obavezna", [])) else ""
                            placeholder = f"{{{{{literatura_ob_key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(literatura_ob))

                            literatura_dop_key = f"literatura_dop_{i}"
                            literatura_dop = data["literatura"].get("dodatna", [])[i - 1] if i - 1 < len(data["literatura"].get("dodatna", [])) else ""
                            placeholder = f"{{{{{literatura_dop_key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(literatura_dop))
    doc.save(output_path)



@app.route('/ask-keywords', methods=['POST'])
def chatKey():
    thread_id = create_thread()

    keys = [
        "naziv_predmeta", "studijski_program", "ects_bodovi", "ciklus",
        "godina_studija", "kod_predmeta", "predavanja", "vjezbe",
        "seminari", "praksa", "vizija", "ishodi_ucenja"
    ]

    for key in keys:
        globals()[key] = request.json.get(key, '')

    user_input = (f"Napravi mi ključne riječi za sljedeći kolegij: Kolegij: {naziv_predmeta}, Studijski program: {studijski_program}, ECTS bodovi: {ects_bodovi}, Ciklus: {ciklus}, Godina studija: {godina_studija}, Kod predmeta: {kod_predmeta}, "
                  f"Predavanja: {predavanja}, Vjezbe: {vjezbe}, Seminari: {seminari}, Praksa: {praksa}, Vizija kolegija (kako je nastavnik zamislio svoje specifičnosti): {vizija}, Početak ishoda učenja studijske grupe: {ishodi_ucenja}. Kraj ishoda učenja studijske grupe.")


    print(user_input)
    submit_message(thread_id, user_input)
    run_id = run_assistant(thread_id)
    status = check_status(run_id, thread_id)
    while status != "completed":
        time.sleep(1)
        status = check_status(run_id, thread_id)

    response = get_response(thread_id)

    print('here is response')
    print(response)

    if response:
        clean_response = response.strip("```json").strip("```").strip()
        data = json.loads(clean_response)
    else:
        return jsonify({"error": "Invalid response from assistant"}), 500

    return jsonify(response)

## I want to return downloadable filled_obrzaca.docx file here
# return send_file(output_path, as_attachment=True)


@app.route('/ask', methods=['POST'])
def chat():
    thread_id = create_thread()

    keys = [
        "naziv_predmeta", "studijski_program", "ects_bodovi", "ciklus",
        "godina_studija", "kod_predmeta", "predavanja", "vjezbe",
        "seminari", "praksa", "vizija", "ishodi_ucenja", "pickedKeywords", "notPickedKeywords"
    ]

    for key in keys:
        globals()[key] = request.json.get(key, '')

    try:
        sati_nastava = int(predavanja) + int(vjezbe) + int(seminari) + int(praksa)
        sati_kolokvij = sati_nastava
        sati_ispit = sati_nastava // 2

        ukupan_broj_sati = sati_nastava + sati_kolokvij + sati_ispit

        ects_nastava = int(ects_bodovi) // 3
        ects_kolokvij = int(ects_bodovi) // 2
        ects_ispit = int(ects_bodovi) - ects_nastava - ects_kolokvij

        user_input = (f"Napravi mi silabus sa sljedeći kolegij: Kolegij: {naziv_predmeta}, Studijski program: {studijski_program}, ECTS bodovi: {ects_bodovi}, Ciklus: {ciklus}, Godina studija: {godina_studija}, Kod predmeta: {kod_predmeta}, "
                    f"Predavanja: {predavanja}, Vjezbe: {vjezbe}, Seminari: {seminari}, Praksa: {praksa}, Ukupan broj sati: {ukupan_broj_sati}, Sati nastave: {sati_nastava}, Sati kolokvij: {sati_kolokvij}, Sati ispit: {sati_ispit}, ECTS nastava: {ects_nastava}, ECTS kolokvij: {ects_kolokvij}, ECTS ispit: {ects_ispit}, Vizija: {vizija}, Ishodi učenja: {ishodi_ucenja}, Ključne riječi: {kljucne_rijeci}")
    except:
        print("Šaljem bez alokacije sati i ects bodova.")
        user_input = (f"Napravi mi silabus sa sljedeći kolegij: Kolegij: {naziv_predmeta}, Studijski program: {studijski_program}, ECTS bodovi: {ects_bodovi}, Ciklus: {ciklus}, Godina studija: {godina_studija}, Kod predmeta: {kod_predmeta}, "
                    f"Predavanja: {predavanja}, Vjezbe: {vjezbe}, Seminari: {seminari}, Praksa: {praksa}, Vizija: {vizija}, Ishodi učenja: {ishodi_ucenja}, Ključne riječi: {kljucne_rijeci}")

    print(user_input)

    submit_message(thread_id, user_input)
    run_id = run_assistant(thread_id)
    status = check_status(run_id, thread_id)
    while status != "completed":
        time.sleep(1)
        status = check_status(run_id, thread_id)

    response = get_response(thread_id)

    print('here is response')
    print(response)

    if response:
        clean_response = response.strip("```json").strip("```").strip()
        data = json.loads(clean_response)
    else:
        return jsonify({"error": "Invalid response from assistant"}), 500

    template_path = os.path.join(app_dir, 'obrazac.docx')

    file_name = f'obrazac_{uuid.uuid4()}.docx' ## generate random string
    output_path = os.path.join(app_dir, file_name)

    for key in keys:
        if key in locals():
            data[key] = locals()[key]

    fill_template(data, template_path, output_path)

    return jsonify({"response": response, "thread_id": thread_id, "document": file_name})


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = '/app/' + filename

    return send_file(
        file_path,
        as_attachment=True,
        mimetype="application/docx",
        download_name=filename
    )

@app.route('/test', methods=['GET'])
def test():
    return jsonify({"message": "Hello, World!"})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8001, debug=True)
